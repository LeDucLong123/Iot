from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "BAO_CAO_DO_AN_IOT_NHA_THONG_MINH_THEO_MAU.docx"
LOGO = Path("assets/report_logo_phenikaa.jpeg")
DIAGRAM_DIR = Path("assets/report_diagrams")

SCHOOL = "ĐẠI HỌC PHENIKAA"
MINISTRY = "BỘ GIÁO DỤC VÀ ĐÀO TẠO"
FACULTY = "KHOA CÔNG NGHỆ THÔNG TIN"
COURSE = "HỆ THỐNG NHÚNG VÀ INTERNET VẠN VẬT"
TITLE = "THIẾT KẾ VÀ MÔ PHỎNG HỆ THỐNG NHÀ THÔNG MINH SỬ DỤNG ESP32, ARDUINO UNO, MQTT VÀ HOME ASSISTANT"
TEACHER = "................................................"
STUDENT = "................................................"
STUDENT_ID = "................................................"
CLASS = "................................................"
MAJOR = "Công nghệ thông tin"
YEAR = "2026"


def set_font(run, size=13, bold=False, italic=False, color="000000"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_page(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.right_margin = Inches(0.98)
    section.left_margin = Inches(1.38)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.27)


def set_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = pg_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            pg_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), "000000")


def clear_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is not None:
        sect_pr.remove(pg_borders)


def p(doc, text="", align=None, before=0, after=6, line=1.3, size=13, bold=False, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return para


def heading(doc, text, level=1):
    if level == 1:
        para = p(doc, text.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=12, line=1.15, size=15, bold=True)
    elif level == 2:
        para = p(doc, text, before=8, after=6, line=1.2, size=13, bold=True)
    else:
        para = p(doc, text, before=6, after=4, line=1.2, size=13, bold=True, italic=True)
    return para


def body(doc, text):
    para = p(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line=1.3, size=13)
    para.paragraph_format.first_line_indent = Inches(0.38)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    set_font(run, size=13)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    set_font(run, size=13)
    return para


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, headers, rows, widths=None, caption=None):
    if caption:
        cap = p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.1, size=12, bold=True)
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    for r_idx, row in enumerate(tbl.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.1
                for run in para.runs:
                    set_font(run, size=12, bold=(r_idx == 0))
    p(doc, "", after=4)
    return tbl


def cover_line(doc, text, size=18, bold=True, after=8):
    return p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, after=after, line=1.0, size=size, bold=bold)


def add_cover(doc, subtitle="ĐỒ ÁN MÔN HỌC", title=TITLE):
    cover_line(doc, MINISTRY, size=17, after=4)
    cover_line(doc, SCHOOL, size=17, after=18)
    if LOGO.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(LOGO), width=Inches(1.75))
        pic.paragraph_format.space_after = Pt(18)
    cover_line(doc, subtitle, size=15, after=3)
    cover_line(doc, COURSE, size=15, after=18)
    cover_line(doc, title, size=18, after=22)
    p(doc, f"Giảng viên hướng dẫn: {TEACHER}", before=0, after=3, size=13.5, bold=True)
    p(doc, f"Sinh viên thực hiện: {STUDENT}", before=0, after=3, size=13.5, bold=True)
    p(doc, f"Mã sinh viên: {STUDENT_ID}", before=0, after=3, size=13.5, bold=True)
    p(doc, f"Lớp: {CLASS}", before=0, after=3, size=13.5, bold=True)
    p(doc, f"Ngành/chuyên ngành: {MAJOR}", before=0, after=3, size=13.5, bold=True)
    for _ in range(2):
        p(doc, "", after=4)
    cover_line(doc, f"Hà Nội, năm {YEAR}", size=14, bold=False, after=0)


def make_diagram(path, title, boxes, arrows):
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1400, 780), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf", 40)
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Times New Roman.ttf", 28)
        font_bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf", 28)
    except Exception:
        font_title = font = font_bold = ImageFont.load_default()
    draw.text((700, 35), title, anchor="mm", fill="black", font=font_title)
    for key, (x, y, w, h, label, fill) in boxes.items():
        draw.rounded_rectangle((x, y, x + w, y + h), radius=14, outline="black", width=3, fill=fill)
        lines = label.split("\\n")
        total = len(lines) * 32
        for idx, line in enumerate(lines):
            draw.text((x + w / 2, y + h / 2 - total / 2 + idx * 32 + 16), line, anchor="mm", fill="black", font=font_bold if idx == 0 else font)
    for a, b in arrows:
        ax, ay, aw, ah, *_ = boxes[a]
        bx, by, bw, bh, *_ = boxes[b]
        start = (ax + aw, ay + ah / 2) if ax < bx else (ax, ay + ah / 2)
        end = (bx, by + bh / 2) if ax < bx else (bx + bw, by + bh / 2)
        draw.line((start, end), fill="black", width=4)
        ex, ey = end
        if ax < bx:
            pts = [(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)]
        else:
            pts = [(ex, ey), (ex + 18, ey - 10), (ex + 18, ey + 10)]
        draw.polygon(pts, fill="black")
    img.save(path)


def create_diagrams():
    block = DIAGRAM_DIR / "so_do_khoi.png"
    flow = DIAGRAM_DIR / "luong_mqtt_i2c.png"
    make_diagram(
        block,
        "Sơ đồ khối hệ thống nhà thông minh",
        {
            "ha": (40, 310, 250, 110, "Home Assistant\\nDashboard", "#EAF3FF"),
            "mqtt": (360, 310, 230, 110, "Mosquitto\\nMQTT Broker", "#E8F8EC"),
            "esp": (660, 285, 250, 160, "ESP32\\nMaster IoT", "#FFF3D6"),
            "uno": (1030, 310, 250, 110, "Arduino Uno\\nI2C Slave", "#FDECEC"),
            "sensor": (650, 560, 270, 110, "Cảm biến\\nRFID, DHT11, mưa, gas", "#F4F4F4"),
            "act": (1010, 560, 300, 110, "Cơ cấu chấp hành\\nservo, LED, buzzer", "#F4F4F4"),
        },
        [("ha", "mqtt"), ("mqtt", "esp"), ("esp", "uno"), ("sensor", "esp"), ("uno", "act")],
    )
    make_diagram(
        flow,
        "Luồng dữ liệu MQTT và I2C",
        {
            "sensor": (50, 310, 230, 110, "Cảm biến\\nđọc dữ liệu", "#F4F4F4"),
            "esp": (350, 285, 260, 160, "ESP32\\nxử lý logic", "#FFF3D6"),
            "mqtt": (700, 310, 230, 110, "MQTT\\npublish/subscribe", "#E8F8EC"),
            "ha": (1010, 310, 270, 110, "Home Assistant\\ngiao diện người dùng", "#EAF3FF"),
            "uno": (350, 560, 260, 110, "Arduino Uno\\nnhận lệnh I2C", "#FDECEC"),
            "act": (700, 560, 250, 110, "Thiết bị\\nthực thi lệnh", "#F4F4F4"),
        },
        [("sensor", "esp"), ("esp", "mqtt"), ("mqtt", "ha"), ("esp", "uno"), ("uno", "act")],
    )
    return block, flow


def build():
    block_img, flow_img = create_diagrams()
    doc = Document()
    for sec in doc.sections:
        set_page(sec)
    set_page_border(doc.sections[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    add_cover(doc)
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(sec)
    set_page_border(sec)
    add_cover(doc, subtitle="ĐỒ ÁN MÔN HỌC:", title=TITLE)

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(sec)
    clear_page_border(sec)
    heading(doc, "LỜI CẢM ƠN", 1)
    body(doc, "Em xin chân thành cảm ơn giảng viên hướng dẫn đã tận tình giảng dạy, định hướng và hỗ trợ em trong quá trình thực hiện đồ án. Những kiến thức về hệ thống nhúng, giao tiếp phần cứng, truyền thông IoT và phương pháp xây dựng báo cáo kỹ thuật là cơ sở quan trọng giúp em hoàn thành đề tài này.")
    body(doc, "Em cũng xin cảm ơn các bạn trong lớp đã hỗ trợ trao đổi ý tưởng, kiểm tra kết nối phần cứng và góp ý trong quá trình thử nghiệm mô hình. Do thời gian thực hiện còn hạn chế, đồ án khó tránh khỏi thiếu sót. Em rất mong nhận được góp ý từ thầy cô để hệ thống được hoàn thiện hơn.")

    doc.add_page_break()
    heading(doc, "LỜI CAM ĐOAN", 1)
    body(doc, "Em xin cam đoan đồ án môn học với đề tài “Thiết kế và mô phỏng hệ thống nhà thông minh sử dụng ESP32, Arduino Uno, MQTT và Home Assistant” là sản phẩm do em tự tìm hiểu, thiết kế, lập trình và tổng hợp báo cáo. Các nội dung tham khảo từ tài liệu kỹ thuật, thư viện phần mềm và tài liệu hướng dẫn đều được ghi nhận trong phần tài liệu tham khảo.")
    body(doc, "Em xin chịu trách nhiệm hoàn toàn về nội dung của đồ án đã nộp.")
    p(doc, "Hà Nội, ngày …… tháng …… năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, before=18, after=18, size=13)
    p(doc, "Sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.RIGHT, after=30, size=13, bold=True)
    p(doc, "(ký và ghi rõ họ tên)", align=WD_ALIGN_PARAGRAPH.RIGHT, after=6, size=13, italic=True)

    doc.add_page_break()
    heading(doc, "PHÂN CÔNG NHIỆM VỤ ĐỒ ÁN", 1)
    table(doc, ["Danh sách các công việc/nhiệm vụ", "Mô tả tóm tắt công việc"], [
        ("Khảo sát yêu cầu", "Xác định chức năng nhà thông minh cần mô phỏng, lựa chọn cảm biến và thiết bị chấp hành."),
        ("Phân tích thiết kế", "Thiết kế kiến trúc ESP32 - Arduino Uno - MQTT - Home Assistant và giao thức I2C."),
        ("Triển khai giải pháp", "Lập trình các module ESP32, Arduino Uno, cấu hình Mosquitto và Home Assistant."),
        ("Kiểm thử, đánh giá", "Kiểm thử từng chức năng, ghi nhận kết quả, đánh giá ưu điểm và hạn chế."),
    ], widths=[2.5, 4.0])
    table(doc, ["Tên sinh viên/Mã sinh viên", "Các công việc", "Tỉ lệ"], [
        (f"{STUDENT} - {STUDENT_ID}", "Khảo sát, thiết kế phần cứng, lập trình ESP32 và Arduino Uno", "40%"),
        ("", "Cấu hình MQTT, Home Assistant, Docker Compose", "25%"),
        ("", "Kiểm thử, đánh giá, viết báo cáo đồ án", "35%"),
    ], widths=[2.2, 3.4, 0.9])

    doc.add_page_break()
    heading(doc, "PHIẾU CHẤM THI TIỂU LUẬN/ĐỒ ÁN", 1)
    p(doc, f"Môn học: {COURSE}                                      Lớp học phần: ................................", size=14)
    p(doc, "Sinh viên thực hiện: ................................................................................", size=14)
    for i in range(1, 6):
        p(doc, f"{i}. .................................................................... Điểm: ........................................", size=14)
    p(doc, "Ngày thi: ............................................................ Phòng thi: ........................................", before=16, size=14)
    p(doc, "", before=50, after=8)
    table(doc, ["Giảng viên chấm thi 1\n(Ký và ghi rõ họ tên)", "Giảng viên chấm thi 2\n(Ký và ghi rõ họ tên)"], [("", "")], widths=[3.2, 3.2])

    doc.add_page_break()
    heading(doc, "MỤC LỤC", 1)
    toc = [
        "LỜI CẢM ƠN",
        "LỜI CAM ĐOAN",
        "PHÂN CÔNG NHIỆM VỤ ĐỒ ÁN",
        "DANH MỤC TỪ VIẾT TẮT",
        "DANH MỤC HÌNH VẼ",
        "DANH MỤC BẢNG",
        "CHƯƠNG 1. TỔNG QUAN VỀ ĐỀ TÀI",
        "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG",
        "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "CHƯƠNG 4. THIẾT KẾ PHẦN CỨNG",
        "CHƯƠNG 5. THIẾT KẾ PHẦN MỀM",
        "CHƯƠNG 6. TRIỂN KHAI MÔ PHỎNG VÀ KIỂM THỬ",
        "CHƯƠNG 7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
        "TÀI LIỆU THAM KHẢO",
    ]
    for i, item in enumerate(toc, 1):
        p(doc, f"{item}{'.' * max(8, 70 - len(item))}{i}", size=13, after=2)

    doc.add_page_break()
    heading(doc, "DANH MỤC TỪ VIẾT TẮT", 1)
    table(doc, ["Chữ viết tắt", "Giải thích"], [
        ("IoT", "Internet of Things - Internet vạn vật"),
        ("ESP32", "Vi điều khiển tích hợp WiFi/Bluetooth của Espressif"),
        ("MQTT", "Message Queuing Telemetry Transport"),
        ("HA", "Home Assistant"),
        ("I2C", "Inter-Integrated Circuit"),
        ("RFID", "Radio Frequency Identification"),
        ("GPIO", "General Purpose Input/Output"),
        ("NTP", "Network Time Protocol"),
    ], widths=[1.7, 4.8])

    doc.add_page_break()
    heading(doc, "DANH MỤC HÌNH VẼ", 1)
    for line in [
        "Hình 3.1 Sơ đồ khối hệ thống nhà thông minh..............................................",
        "Hình 3.2 Luồng dữ liệu MQTT và I2C........................................................",
    ]:
        p(doc, line, size=13, after=2)
    heading(doc, "DANH MỤC BẢNG", 1)
    for line in [
        "Bảng 4.1 Danh sách linh kiện sử dụng......................................................",
        "Bảng 4.2 Kết nối thiết bị với ESP32........................................................",
        "Bảng 4.3 Kết nối thiết bị với Arduino Uno..................................................",
        "Bảng 5.1 Giao thức điều khiển I2C.........................................................",
        "Bảng 5.2 Các topic MQTT chính................................................................",
        "Bảng 6.1 Kết quả kiểm thử chức năng.......................................................",
    ]:
        p(doc, line, size=13, after=2)

    doc.add_page_break()
    heading(doc, "CHƯƠNG 1. TỔNG QUAN VỀ ĐỀ TÀI", 1)
    heading(doc, "1.1. Lời mở đầu", 2)
    body(doc, "Ngày nay, Internet of Things đã trở thành một hướng phát triển quan trọng trong lĩnh vực tự động hóa và hệ thống nhúng. Các thiết bị trong gia đình không chỉ được điều khiển trực tiếp bằng công tắc mà còn có thể được giám sát, điều khiển từ xa và tự động hóa theo dữ liệu cảm biến.")
    body(doc, "Trong bối cảnh đó, việc xây dựng một mô hình nhà thông minh giúp sinh viên hiểu rõ cách phối hợp giữa vi điều khiển, cảm biến, giao thức truyền thông và giao diện điều khiển. Đề tài sử dụng ESP32 làm bộ điều khiển trung tâm, Arduino Uno làm bộ mở rộng ngõ ra, MQTT làm giao thức truyền dữ liệu và Home Assistant làm nền tảng giám sát.")
    heading(doc, "1.2. Lý do chọn đề tài", 2)
    body(doc, "ESP32 có WiFi tích hợp, giá thành thấp và hỗ trợ nhiều chuẩn giao tiếp như SPI, I2C, UART, PWM nên rất phù hợp với các ứng dụng IoT. Khi số lượng thiết bị mô phỏng trong nhà tăng lên, việc dùng thêm Arduino Uno làm I2C slave giúp mở rộng số ngõ ra, đồng thời minh họa mô hình hệ thống nhúng nhiều vi điều khiển.")
    body(doc, "Đề tài có tính thực tiễn vì tích hợp nhiều tình huống thường gặp trong nhà thông minh như điều khiển đèn, khóa cửa RFID, đo nhiệt độ - độ ẩm, tự động bật quạt, che mưa, phát hiện người trong nhà vệ sinh và cảnh báo an toàn khi có gas hoặc rung chấn.")
    heading(doc, "1.3. Mục tiêu đề tài", 2)
    for item in [
        "Xây dựng mô hình nhà thông minh có khả năng điều khiển và giám sát qua Home Assistant.",
        "Ứng dụng ESP32 để kết nối WiFi, MQTT, đọc cảm biến và xử lý logic trung tâm.",
        "Ứng dụng Arduino Uno làm thiết bị slave nhận lệnh I2C để điều khiển servo và LED.",
        "Tích hợp các cảm biến RFID, DHT11, cảm biến mưa LM393, HC-SR04, cảm biến rung và cảm biến gas.",
        "Đánh giá hoạt động của hệ thống thông qua các kịch bản mô phỏng và bảng kiểm thử.",
    ]:
        bullet(doc, item)
    heading(doc, "1.4. Phạm vi thực hiện", 2)
    body(doc, "Đồ án tập trung vào mô hình prototype ở quy mô phòng thí nghiệm. Các LED được sử dụng để mô phỏng đèn trong nhà, servo mô phỏng cửa chính và cửa sổ/chan mưa, buzzer mô phỏng cảnh báo khẩn cấp, còn quạt được mô phỏng bằng tải điều khiển qua GPIO12 hoặc module L298N.")

    doc.add_page_break()
    heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    for title, text in [
        ("2.1. Internet of Things", "IoT là mô hình kết nối các thiết bị vật lý với mạng để thu thập dữ liệu, trao đổi thông tin và thực hiện điều khiển. Một hệ thống IoT thường gồm cảm biến, bộ xử lý nhúng, giao thức truyền thông, máy chủ hoặc broker và giao diện người dùng."),
        ("2.2. Vi điều khiển ESP32", "ESP32 là vi điều khiển tích hợp WiFi/Bluetooth, có nhiều chân GPIO và hỗ trợ các giao tiếp ngoại vi như SPI, I2C, UART, PWM. Trong đề tài, ESP32 đảm nhiệm kết nối WiFi, MQTT, đọc cảm biến, xử lý logic và gửi lệnh I2C."),
        ("2.3. Arduino Uno", "Arduino Uno sử dụng vi điều khiển ATmega328P, phù hợp để điều khiển LED, servo và các cơ cấu đơn giản. Trong hệ thống, Arduino Uno đóng vai trò I2C slave tại địa chỉ 0x08."),
        ("2.4. Giao thức MQTT", "MQTT là giao thức publish/subscribe nhẹ, thường dùng trong IoT. ESP32 publish dữ liệu cảm biến lên broker Mosquitto, Home Assistant subscribe để hiển thị; khi người dùng thao tác, Home Assistant publish lệnh để ESP32 nhận và xử lý."),
        ("2.5. Home Assistant", "Home Assistant là nền tảng nhà thông minh mã nguồn mở, cho phép tạo dashboard, quản lý entity MQTT và theo dõi trạng thái thiết bị. Trong đồ án, các switch, sensor và binary sensor được khai báo trong configuration.yaml."),
        ("2.6. Giao tiếp I2C", "I2C sử dụng hai đường SDA và SCL. ESP32 dùng GPIO21 làm SDA và GPIO22 làm SCL, kết nối đến A4/A5 của Arduino Uno. Do ESP32 dùng mức logic 3.3V, hai đường I2C được kéo lên 3.3V bằng điện trở 4.7kΩ hoặc 10kΩ."),
    ]:
        heading(doc, title, 2)
        body(doc, text)

    doc.add_page_break()
    heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    heading(doc, "3.1. Yêu cầu chức năng", 2)
    for item in [
        "Điều khiển đèn phòng khách, bếp, phòng ngủ, nhà vệ sinh và LED trạng thái.",
        "Mở cửa bằng thẻ RFID hợp lệ, cảnh báo khi thẻ không hợp lệ và tự đóng cửa sau 3 giây.",
        "Đọc nhiệt độ, độ ẩm và tự động bật quạt khi nhiệt độ vượt ngưỡng 30.0°C.",
        "Phát hiện mưa để tự động điều khiển cơ cấu cửa sổ/chan mưa.",
        "Phát hiện người trong nhà vệ sinh bằng HC-SR04 để tự động bật/tắt đèn.",
        "Cảnh báo khi phát hiện rung chấn hoặc rò rỉ khí gas.",
        "Hiển thị ngày giờ, nhiệt độ và độ ẩm trên LCD I2C.",
    ]:
        bullet(doc, item)
    heading(doc, "3.2. Kiến trúc tổng thể", 2)
    body(doc, "Hệ thống được chia thành ba lớp: lớp thiết bị hiện trường, lớp điều khiển nhúng và lớp giám sát - điều khiển. Lớp thiết bị gồm cảm biến và cơ cấu chấp hành. Lớp điều khiển nhúng gồm ESP32 và Arduino Uno. Lớp giám sát gồm Mosquitto và Home Assistant.")
    doc.add_picture(str(block_img), width=Inches(5.8))
    p(doc, "Hình 3.1 Sơ đồ khối hệ thống nhà thông minh", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    heading(doc, "3.3. Luồng dữ liệu", 2)
    body(doc, "Dữ liệu cảm biến được ESP32 đọc và xử lý định kỳ. Khi có dữ liệu mới hoặc trạng thái thay đổi, ESP32 publish lên MQTT Broker. Home Assistant hiển thị trạng thái và cho phép người dùng gửi lệnh điều khiển. Nếu thiết bị cần điều khiển nằm trên Arduino Uno, ESP32 chuyển lệnh MQTT thành gói I2C gồm mã lệnh và dữ liệu.")
    doc.add_picture(str(flow_img), width=Inches(5.8))
    p(doc, "Hình 3.2 Luồng dữ liệu MQTT và I2C", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)

    doc.add_page_break()
    heading(doc, "CHƯƠNG 4. THIẾT KẾ PHẦN CỨNG", 1)
    table(doc, ["STT", "Linh kiện", "Chức năng"], [
        (1, "ESP32 DevKit", "Điều khiển trung tâm, WiFi, MQTT, đọc cảm biến."),
        (2, "Arduino Uno", "I2C slave, mở rộng ngõ ra."),
        (3, "RFID MFRC522", "Đọc UID thẻ mở cửa."),
        (4, "DHT11", "Đo nhiệt độ và độ ẩm."),
        (5, "LM393 Rain Sensor", "Phát hiện mưa."),
        (6, "HC-SR04", "Phát hiện người bằng khoảng cách."),
        (7, "Cảm biến rung", "Cảnh báo động đất/rung chấn."),
        (8, "Cảm biến gas MQ", "Cảnh báo rò rỉ khí gas."),
        (9, "LCD I2C 16x2", "Hiển thị thời gian và thông số môi trường."),
        (10, "Servo, LED, buzzer, L298N", "Mô phỏng cơ cấu chấp hành và cảnh báo."),
    ], widths=[0.45, 2.0, 4.0], caption="Bảng 4.1 Danh sách linh kiện sử dụng")
    table(doc, ["Thiết bị", "Chân ESP32", "Ghi chú"], [
        ("I2C SDA/SCL", "GPIO21/GPIO22", "Kết nối Arduino A4/A5 và LCD I2C, kéo lên 3.3V."),
        ("RFID MFRC522", "5/18/23/19/14", "SS/SCK/MOSI/MISO/RST qua SPI."),
        ("DHT11", "GPIO15", "Đọc nhiệt độ và độ ẩm."),
        ("Cảm biến mưa", "GPIO27", "LOW khi có mưa."),
        ("HC-SR04", "GPIO25/GPIO26", "TRIG/ECHO."),
        ("Cảm biến rung", "GPIO32", "HIGH khi có rung."),
        ("Cảm biến gas", "GPIO4", "LOW khi phát hiện gas."),
        ("Buzzer", "GPIO33", "Cảnh báo gas hoặc động đất."),
        ("Quạt/L298N", "GPIO12", "Điều khiển quạt trong code hiện tại."),
    ], widths=[1.8, 1.5, 3.2], caption="Bảng 4.2 Kết nối thiết bị với ESP32")
    table(doc, ["Thiết bị", "Chân Arduino", "Chức năng"], [
        ("Servo cửa chính", "D12", "Mở/đóng cửa chính."),
        ("Servo cửa sổ/chan mưa", "D11", "Điều khiển cơ cấu che mưa."),
        ("LED phòng khách", "D7", "Mô phỏng đèn phòng khách."),
        ("LED RFID xanh/đỏ", "D8/D9", "Báo thẻ hợp lệ hoặc không hợp lệ."),
        ("Đèn bếp", "D6", "Mô phỏng đèn bếp hoặc cảnh báo lỗi RFID theo cách đấu."),
        ("Đèn phòng ngủ", "D5", "Mô phỏng đèn phòng ngủ."),
        ("Đèn nhà vệ sinh", "D4", "Mô phỏng đèn nhà vệ sinh."),
    ], widths=[2.2, 1.3, 3.0], caption="Bảng 4.3 Kết nối thiết bị với Arduino Uno")
    heading(doc, "4.4. Lưu ý đấu nối", 2)
    body(doc, "ESP32 và Arduino Uno bắt buộc phải nối chung GND. Module RFID nên cấp nguồn 3.3V. Servo và tải công suất như quạt cần nguồn riêng ổn định, không cấp trực tiếp từ chân GPIO. Đường I2C nên ngắn và có điện trở kéo lên 3.3V để tránh lỗi truyền thông.")

    doc.add_page_break()
    heading(doc, "CHƯƠNG 5. THIẾT KẾ PHẦN MỀM", 1)
    heading(doc, "5.1. Cấu trúc chương trình", 2)
    body(doc, "Dự án sử dụng PlatformIO với hai môi trường build: esp32 và uno. Mã nguồn được tách thành hai nhánh trong thư mục src/esp32 và src/arduino. Các chức năng trên ESP32 được tách thành module như mqtt_manager, dht_manager, rain_manager, rfid_servo_manager, radar_manager, gas_manager và earthquake_manager.")
    heading(doc, "5.2. Giao thức I2C", 2)
    table(doc, ["Command ID", "Dữ liệu", "Tác dụng"], [
        ("0x01", "90 / 0", "Mở hoặc đóng servo cửa chính."),
        ("0x02", "1", "Kích hoạt LED đỏ/còi báo thẻ RFID sai."),
        ("0x03", "1 / 0", "Bật/tắt đèn phòng khách."),
        ("0x05", "0 / 90", "Điều khiển servo cửa sổ/chan mưa."),
        ("0x06", "1 / 0", "Bật/tắt đèn bếp."),
        ("0x07", "1 / 0", "Bật/tắt đèn phòng ngủ."),
        ("0x08", "1 / 0", "Bật/tắt đèn nhà vệ sinh."),
    ], widths=[1.25, 1.25, 4.0], caption="Bảng 5.1 Giao thức điều khiển I2C")
    doc.add_page_break()
    heading(doc, "5.3. Hệ thống topic MQTT", 2)
    table(doc, ["Topic", "Payload", "Chức năng"], [
        ("home/esp32/led1/set", "ON/OFF", "Điều khiển LED 1."),
        ("home/esp32/servo/set", "OPEN/CLOSE", "Điều khiển cửa chính."),
        ("home/esp32/fan/set", "ON/OFF", "Điều khiển quạt thủ công."),
        ("home/esp32/fan/mode/set", "auto/manual", "Chọn chế độ quạt."),
        ("home/esp32/window/set", "OPEN/CLOSE", "Điều khiển cửa sổ/chan mưa."),
        ("home/esp32/sensor/temperature", "Số thực", "Publish nhiệt độ."),
        ("home/esp32/sensor/humidity", "Số thực", "Publish độ ẩm."),
        ("home/esp32/sensor/rain", "ON/OFF", "Trạng thái mưa."),
        ("home/esp32/sensor/gas", "ON/OFF", "Cảnh báo gas."),
        ("home/esp32/sensor/earthquake", "ON/OFF", "Cảnh báo rung chấn."),
    ], widths=[2.8, 1.25, 2.45], caption="Bảng 5.2 Các topic MQTT chính")
    heading(doc, "5.4. Logic các chức năng chính", 2)
    body(doc, "Với khóa cửa RFID, ESP32 đọc UID từ MFRC522 và so sánh với danh sách thẻ hợp lệ trong config.cpp. Nếu hợp lệ, ESP32 gửi lệnh I2C mở cửa và tự đóng sau 3 giây; nếu không hợp lệ, hệ thống kích hoạt cảnh báo lỗi.")
    body(doc, "Với DHT11, hệ thống đọc dữ liệu mỗi 5 giây và publish nhiệt độ, độ ẩm lên MQTT. Khi chế độ quạt tự động được bật, nếu nhiệt độ vượt 30.0°C, ESP32 bật quạt qua GPIO12; nếu nhiệt độ giảm, quạt tắt.")
    body(doc, "Với cảm biến mưa, ESP32 kiểm tra chân DO mỗi 500 ms. Khi có mưa, hệ thống điều khiển servo cửa sổ/chan mưa qua I2C và đồng bộ trạng thái lên Home Assistant. Khi hết mưa, hệ thống chờ 5 giây trước khi đưa cơ cấu về trạng thái bình thường.")
    body(doc, "Với HC-SR04, nếu phát hiện vật thể ở khoảng cách nhỏ hơn 10 cm, hệ thống bật đèn nhà vệ sinh; khi không còn phát hiện người, hệ thống chờ 5 giây rồi tắt đèn. Với cảm biến gas và rung, trạng thái cảnh báo được publish lên MQTT và buzzer GPIO33 được kích hoạt.")

    doc.add_page_break()
    heading(doc, "CHƯƠNG 6. TRIỂN KHAI MÔ PHỎNG VÀ KIỂM THỬ", 1)
    heading(doc, "6.1. Triển khai server cục bộ", 2)
    body(doc, "Mosquitto và Home Assistant được triển khai bằng Docker Compose. Mosquitto mở cổng 1883 để ESP32 kết nối MQTT. Home Assistant mở cổng 8123 để người dùng truy cập dashboard. Khi triển khai, ESP32 và máy chạy Docker cần nằm trong cùng mạng LAN.")
    heading(doc, "6.2. Quy trình chạy thử", 2)
    for item in [
        "Chạy server bằng lệnh docker compose up -d.",
        "Biên dịch và nạp chương trình ESP32 bằng PlatformIO với môi trường esp32.",
        "Biên dịch và nạp chương trình Arduino Uno với môi trường uno.",
        "Khởi động Serial Monitor để kiểm tra WiFi, MQTT và I2C.",
        "Mở Home Assistant, kiểm tra các switch, sensor và binary sensor MQTT.",
    ]:
        numbered(doc, item)
    heading(doc, "6.3. Kết quả kiểm thử", 2)
    table(doc, ["STT", "Chức năng", "Kết quả mong đợi", "Đánh giá"], [
        (1, "Kết nối WiFi/MQTT", "ESP32 nhận IP và subscribe topic thành công.", "Đạt"),
        (2, "Điều khiển đèn", "LED/đèn mô phỏng thay đổi theo Home Assistant.", "Đạt"),
        (3, "RFID hợp lệ", "Servo mở cửa, LED xanh sáng, publish OPEN.", "Đạt"),
        (4, "RFID sai", "LED đỏ/cảnh báo lỗi, publish unauthorized.", "Đạt"),
        (5, "DHT11 và quạt", "Publish nhiệt độ/độ ẩm, quạt bật khi > 30.0°C.", "Đạt"),
        (6, "Cảm biến mưa", "Rain ON và servo chan mưa hoạt động.", "Đạt"),
        (7, "HC-SR04", "Đèn nhà vệ sinh bật khi phát hiện người.", "Đạt"),
        (8, "Gas/rung", "Binary sensor ON và buzzer kích hoạt.", "Đạt"),
        (9, "LCD", "Hiển thị ngày giờ, nhiệt độ và độ ẩm.", "Đạt"),
    ], widths=[0.45, 1.7, 3.4, 0.9], caption="Bảng 6.1 Kết quả kiểm thử chức năng")
    heading(doc, "6.4. Đánh giá", 2)
    body(doc, "Kết quả kiểm thử cho thấy hệ thống đáp ứng được các chức năng chính của mô hình nhà thông minh. Việc tách module phần mềm giúp chương trình dễ đọc, dễ bảo trì. MQTT giúp đồng bộ trạng thái giữa ESP32 và Home Assistant tương đối nhanh, trong khi I2C giúp mở rộng ngõ ra qua Arduino Uno.")
    body(doc, "Hạn chế của mô hình là MQTT Broker đang cho phép anonymous để tiện mô phỏng, danh sách UID RFID còn cố định trong code và các cảm biến như DHT11, cảm biến gas cần hiệu chuẩn nếu đưa vào sử dụng thực tế.")

    doc.add_page_break()
    heading(doc, "CHƯƠNG 7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    heading(doc, "7.1. Kết luận", 2)
    body(doc, "Đồ án đã thiết kế và mô phỏng thành công hệ thống nhà thông minh sử dụng ESP32, Arduino Uno, MQTT và Home Assistant. Hệ thống thực hiện được các chức năng điều khiển đèn, khóa cửa RFID, đo nhiệt độ - độ ẩm, điều khiển quạt tự động, xử lý khi có mưa, tự động bật đèn nhà vệ sinh và cảnh báo an toàn khi có gas hoặc rung chấn.")
    body(doc, "Thông qua đề tài, người thực hiện nắm được cách kết hợp phần cứng nhúng, giao tiếp I2C, giao thức MQTT, Docker và Home Assistant để tạo thành một hệ thống IoT hoàn chỉnh ở mức prototype.")
    heading(doc, "7.2. Hướng phát triển", 2)
    for item in [
        "Bổ sung username/password hoặc TLS cho MQTT Broker.",
        "Xây dựng giao diện thêm/xóa thẻ RFID từ Home Assistant.",
        "Lưu dữ liệu cảm biến vào InfluxDB và hiển thị biểu đồ bằng Grafana.",
        "Gửi cảnh báo qua Telegram, email hoặc push notification.",
        "Thiết kế PCB, vỏ bảo vệ và nguồn riêng cho servo, quạt.",
        "Nâng cấp DHT11 lên DHT22 hoặc SHT31 để tăng độ chính xác.",
    ]:
        bullet(doc, item)

    doc.add_page_break()
    heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    for item in [
        "Espressif Systems, ESP32 Technical Reference Manual.",
        "Arduino Documentation, Arduino Uno Rev3.",
        "MQTT Version 3.1.1 Specification.",
        "Home Assistant Documentation - MQTT Integration.",
        "Eclipse Mosquitto Documentation.",
        "PlatformIO Documentation.",
        "Datasheet DHT11, MFRC522, HC-SR04 và các module cảm biến sử dụng trong đồ án.",
        "Tài liệu hướng dẫn đấu nối và cấu hình trong thư mục .doc của dự án.",
    ]:
        numbered(doc, item)

    doc.save(OUT)


if __name__ == "__main__":
    build()

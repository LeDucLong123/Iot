from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "BAO_CAO_DO_AN_IOT_NHA_THONG_MINH.docx"


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_table(table, header=True, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if widths:
        set_table_width(table, widths)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_font(run, size=9.5)
            if header and r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_paragraph(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    color = "2E74B5" if level in (1, 2) else "1F4D78"
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=color)
    p.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}.get(level, 6))
    p.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}.get(level, 4))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r)


def add_kv_table(doc, rows, widths=(1.9, 4.6)):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Hạng mục"
    table.rows[0].cells[1].text = "Nội dung"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    style_table(table, widths=widths)
    return table


def add_data_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    style_table(table, widths=widths)
    return table


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Trang ")
    set_font(r, size=9, color="666666")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r._r.append(fld_char1)
    r._r.append(instr_text)
    r._r.append(fld_char2)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    add_page_number(section)

    # Cover page - editorial cover pattern for a formal technical report.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRƯỜNG/VIỆN: ................................................")
    set_font(r, size=12, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KHOA/BỘ MÔN: ................................................")
    set_font(r, size=12, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO ĐỒ ÁN")
    set_font(r, size=22, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THIẾT KẾ VÀ MÔ PHỎNG HỆ THỐNG NHÀ THÔNG MINH SỬ DỤNG ESP32, ARDUINO UNO, MQTT VÀ HOME ASSISTANT")
    set_font(r, size=16, bold=True, color="2E74B5")

    doc.add_paragraph()
    add_kv_table(doc, [
        ("Sinh viên thực hiện", "................................................"),
        ("Mã số sinh viên", "................................................"),
        ("Lớp", "................................................"),
        ("Giảng viên hướng dẫn", "................................................"),
        ("Đơn vị", "................................................"),
        ("Năm học", "2025 - 2026"),
    ], widths=(2.1, 4.2))

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TP. Hồ Chí Minh, tháng 06 năm 2026")
    set_font(r, size=11, italic=True)

    doc.add_page_break()

    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_paragraph(doc, "Trong quá trình thực hiện đồ án, em đã có cơ hội vận dụng các kiến thức về hệ thống nhúng, lập trình vi điều khiển, giao tiếp I2C, truyền thông MQTT và nền tảng Home Assistant để xây dựng một mô hình nhà thông minh hoàn chỉnh ở mức prototype.")
    add_paragraph(doc, "Em xin gửi lời cảm ơn đến giảng viên hướng dẫn đã định hướng, góp ý và hỗ trợ em trong quá trình triển khai đề tài. Em cũng xin cảm ơn các bạn trong lớp đã hỗ trợ trao đổi ý tưởng, kiểm tra kết nối phần cứng và thử nghiệm các chức năng của hệ thống.")
    add_paragraph(doc, "Do thời gian thực hiện và điều kiện thiết bị còn hạn chế, báo cáo khó tránh khỏi thiếu sót. Em rất mong nhận được ý kiến đóng góp của thầy cô để đề tài được hoàn thiện hơn.")

    add_heading(doc, "TÓM TẮT ĐỒ ÁN", 1)
    add_paragraph(doc, "Đồ án trình bày quá trình thiết kế và mô phỏng hệ thống nhà thông minh sử dụng ESP32 làm bộ điều khiển trung tâm, Arduino Uno làm bộ mở rộng ngõ ra, MQTT làm giao thức truyền thông và Home Assistant làm giao diện giám sát - điều khiển. Hệ thống cho phép người dùng điều khiển đèn, cửa, cửa sổ, quạt làm mát; đồng thời tự động phản ứng với các tình huống môi trường như nhiệt độ cao, trời mưa, có người trong nhà vệ sinh, rung chấn và rò rỉ khí gas.")
    add_paragraph(doc, "Mô hình được xây dựng theo kiến trúc IoT cục bộ. ESP32 kết nối WiFi, trao đổi dữ liệu với MQTT Broker Mosquitto, đồng bộ trạng thái với Home Assistant và gửi lệnh I2C đến Arduino Uno. Arduino Uno tiếp nhận các gói lệnh I2C để điều khiển servo cửa chính, servo cửa sổ và các đèn LED mô phỏng khu vực trong nhà.")
    add_paragraph(doc, "Kết quả đạt được là một hệ thống mô phỏng có đầy đủ các chức năng cơ bản của nhà thông minh: điều khiển từ xa, đồng bộ trạng thái bằng MQTT, tự động hóa theo cảm biến, cảnh báo an toàn và hiển thị thông tin trên LCD.")
    add_paragraph(doc, "Từ khóa: ESP32, Arduino Uno, IoT, MQTT, Home Assistant, Mosquitto, I2C, RFID, DHT11, nhà thông minh.")

    add_heading(doc, "MỤC LỤC", 1)
    add_numbered(doc, [
        "Giới thiệu đề tài",
        "Cơ sở lý thuyết",
        "Phân tích yêu cầu hệ thống",
        "Thiết kế tổng thể",
        "Thiết kế phần cứng",
        "Thiết kế phần mềm và giao thức",
        "Thiết kế các chức năng chính",
        "Mô phỏng, cài đặt và vận hành",
        "Kiểm thử và đánh giá kết quả",
        "Kết luận và hướng phát triển",
        "Tài liệu tham khảo và phụ lục",
    ])

    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_paragraph(doc, "Internet of Things đang trở thành nền tảng quan trọng trong các hệ thống tự động hóa hiện đại. Trong lĩnh vực nhà thông minh, các thiết bị điện không chỉ được điều khiển bằng công tắc truyền thống mà còn có thể được giám sát, điều khiển từ xa và tự động hóa theo dữ liệu cảm biến.")
    add_paragraph(doc, "ESP32 là vi điều khiển phù hợp cho các ứng dụng IoT vì có WiFi tích hợp, nhiều chân GPIO, hỗ trợ SPI/I2C/UART/PWM và có thể lập trình thuận tiện bằng Arduino Framework. Khi số lượng thiết bị ngoại vi tăng, việc sử dụng thêm Arduino Uno làm slave I2C giúp mở rộng ngõ ra và minh họa cách phối hợp nhiều vi điều khiển trong một hệ thống nhúng.")
    add_paragraph(doc, "Vì vậy, đề tài được lựa chọn nhằm xây dựng một mô hình nhà thông minh có tính thực hành cao, bám sát các công nghệ phổ biến trong IoT như MQTT, Home Assistant và Docker.")

    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    add_bullets(doc, [
        "Xây dựng mô hình nhà thông minh có khả năng giám sát và điều khiển qua Home Assistant.",
        "Ứng dụng ESP32 làm bộ điều khiển trung tâm, kết nối WiFi và MQTT.",
        "Ứng dụng Arduino Uno làm thiết bị slave nhận lệnh I2C để điều khiển servo và LED mô phỏng.",
        "Tích hợp RFID, DHT11, cảm biến mưa, HC-SR04, cảm biến rung và cảm biến gas.",
        "Xây dựng các kịch bản tự động hóa: mở cửa bằng RFID, tự động che mưa, tự động bật quạt, tự động bật đèn nhà vệ sinh, cảnh báo gas và động đất.",
    ])

    add_heading(doc, "1.3. Phạm vi thực hiện", 2)
    add_paragraph(doc, "Đồ án tập trung vào mô phỏng và xây dựng prototype quy mô nhỏ. Servo được dùng để mô phỏng cửa chính và cửa sổ/chan mưa, LED được dùng để mô phỏng đèn trong nhà, buzzer được dùng để mô phỏng cảnh báo khẩn cấp. Hệ thống chưa đi sâu vào thiết kế PCB, vỏ hộp, bảo mật mạng nâng cao hoặc triển khai cloud public.")

    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", 1)
    add_heading(doc, "2.1. Kiến trúc IoT", 2)
    add_paragraph(doc, "Một hệ thống IoT cơ bản gồm thiết bị cảm biến, bộ xử lý nhúng, giao thức truyền thông, nền tảng quản lý và giao diện người dùng. Trong đồ án này, ESP32 là node IoT trung tâm, Mosquitto là MQTT Broker, Home Assistant là nền tảng giám sát - điều khiển và Arduino Uno là bộ mở rộng ngõ ra.")
    add_heading(doc, "2.2. Giao thức I2C", 2)
    add_paragraph(doc, "I2C là giao thức nối tiếp đồng bộ gồm hai đường SDA và SCL. ESP32 đóng vai trò master, Arduino Uno đóng vai trò slave tại địa chỉ 0x08. Do ESP32 dùng mức logic 3.3V còn Arduino Uno dùng 5V, hai đường SDA/SCL được kéo lên 3.3V bằng điện trở 4.7kΩ hoặc 10kΩ để tăng độ an toàn khi kết nối trực tiếp.")
    add_heading(doc, "2.3. MQTT và Home Assistant", 2)
    add_paragraph(doc, "MQTT hoạt động theo mô hình publish/subscribe, phù hợp với hệ thống IoT vì nhẹ, dễ triển khai và phản hồi nhanh trong mạng LAN. Home Assistant subscribe các topic trạng thái để hiển thị dữ liệu, đồng thời publish lệnh điều khiển khi người dùng thao tác trên giao diện.")

    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG", 1)
    add_data_table(doc, ["Nhóm yêu cầu", "Mô tả"], [
        ("Điều khiển thiết bị", "Bật/tắt đèn, mở/đóng cửa, điều khiển quạt và cửa sổ qua Home Assistant."),
        ("Tự động hóa", "Tự bật quạt khi nhiệt độ cao, tự xử lý khi mưa, tự bật đèn nhà vệ sinh khi có người."),
        ("An toàn", "Cảnh báo khi có rung chấn hoặc rò rỉ khí gas bằng MQTT và buzzer."),
        ("Hiển thị", "LCD I2C hiển thị ngày giờ, nhiệt độ và độ ẩm."),
        ("Mở rộng", "Tách module phần mềm rõ ràng, dùng I2C command ID để dễ thêm thiết bị."),
    ], widths=(1.8, 4.7))

    add_heading(doc, "CHƯƠNG 4. THIẾT KẾ TỔNG THỂ", 1)
    add_paragraph(doc, "Hệ thống được thiết kế theo ba lớp: lớp thiết bị hiện trường, lớp điều khiển nhúng và lớp giám sát - điều khiển. Lớp thiết bị gồm cảm biến và cơ cấu chấp hành. Lớp điều khiển nhúng gồm ESP32 và Arduino Uno. Lớp giám sát gồm Mosquitto và Home Assistant chạy bằng Docker.")
    add_data_table(doc, ["Thành phần", "Vai trò"], [
        ("ESP32", "Kết nối WiFi, MQTT, đọc cảm biến, xử lý logic chính, gửi lệnh I2C."),
        ("Arduino Uno", "Nhận lệnh I2C từ ESP32 và điều khiển servo, LED mô phỏng."),
        ("Mosquitto", "MQTT Broker trung gian giữa ESP32 và Home Assistant."),
        ("Home Assistant", "Giao diện điều khiển, giám sát, đồng bộ trạng thái thiết bị."),
        ("Docker Compose", "Triển khai Mosquitto và Home Assistant trong môi trường cục bộ."),
    ], widths=(1.6, 4.9))
    add_paragraph(doc, "Luồng dữ liệu tổng quát: cảm biến gửi dữ liệu về ESP32, ESP32 publish lên MQTT, Home Assistant hiển thị dữ liệu. Khi người dùng điều khiển trên Home Assistant, lệnh được publish vào topic MQTT, ESP32 nhận lệnh, xử lý và nếu cần sẽ gửi tiếp lệnh I2C xuống Arduino Uno.")

    add_heading(doc, "CHƯƠNG 5. THIẾT KẾ PHẦN CỨNG", 1)
    add_heading(doc, "5.1. Danh sách linh kiện", 2)
    add_data_table(doc, ["STT", "Linh kiện", "Chức năng"], [
        (1, "ESP32 DevKit", "Điều khiển trung tâm, WiFi, MQTT, đọc cảm biến."),
        (2, "Arduino Uno", "Slave I2C, mở rộng ngõ ra."),
        (3, "MFRC522 RFID", "Đọc thẻ mở cửa."),
        (4, "DHT11", "Đo nhiệt độ và độ ẩm."),
        (5, "LM393 Rain Sensor", "Phát hiện mưa."),
        (6, "HC-SR04", "Phát hiện người bằng khoảng cách."),
        (7, "Cảm biến rung", "Cảnh báo động đất/rung chấn."),
        (8, "Cảm biến gas MQ", "Cảnh báo rò rỉ khí gas."),
        (9, "LCD I2C 16x2", "Hiển thị ngày giờ, nhiệt độ, độ ẩm."),
        (10, "Servo SG90, LED, buzzer, L298N", "Mô phỏng cơ cấu chấp hành và cảnh báo."),
    ], widths=(0.55, 2.0, 3.95))

    doc.add_page_break()
    add_heading(doc, "5.2. Bảng kết nối ESP32", 2)
    add_data_table(doc, ["Thiết bị", "Chân ESP32", "Ghi chú"], [
        ("I2C SDA", "GPIO21", "Kết nối Arduino A4 và LCD SDA, kéo lên 3.3V."),
        ("I2C SCL", "GPIO22", "Kết nối Arduino A5 và LCD SCL, kéo lên 3.3V."),
        ("RFID SS/SCK/MOSI/MISO/RST", "5/18/23/19/14", "Giao tiếp SPI với MFRC522."),
        ("DHT11", "GPIO15", "Đọc nhiệt độ và độ ẩm."),
        ("Cảm biến mưa", "GPIO27", "LOW khi có mưa."),
        ("HC-SR04", "GPIO25/GPIO26", "TRIG/ECHO."),
        ("Cảm biến rung", "GPIO32", "HIGH khi có rung."),
        ("Cảm biến gas", "GPIO4", "LOW khi phát hiện gas."),
        ("Buzzer", "GPIO33", "Cảnh báo gas hoặc động đất."),
        ("Quạt/L298N", "GPIO12", "Điều khiển quạt trong code hiện tại."),
    ], widths=(2.2, 1.5, 2.8))

    doc.add_page_break()
    add_heading(doc, "5.3. Bảng kết nối Arduino Uno", 2)
    add_data_table(doc, ["Thiết bị", "Chân Arduino", "Chức năng"], [
        ("Servo cửa chính", "D12", "Mở/đóng cửa chính."),
        ("Servo cửa sổ/chan mưa", "D11", "Điều khiển cơ cấu che mưa."),
        ("LED phòng khách", "D7", "Mô phỏng đèn phòng khách."),
        ("LED RFID xanh", "D8", "Báo thẻ hợp lệ."),
        ("LED RFID đỏ", "D9", "Báo thẻ không hợp lệ."),
        ("Đèn bếp/cảnh báo lỗi RFID", "D6", "Ngõ ra mô phỏng theo code hiện tại."),
        ("Đèn phòng ngủ", "D5", "Mô phỏng đèn phòng ngủ."),
        ("Đèn nhà vệ sinh", "D4", "Mô phỏng đèn nhà vệ sinh."),
    ], widths=(2.4, 1.4, 2.7))
    add_paragraph(doc, "Lưu ý nguồn: tất cả board phải nối chung GND; RFID dùng 3.3V; servo nên dùng nguồn 5V ổn định; không cấp tải công suất như quạt trực tiếp từ GPIO mà cần relay, transistor hoặc L298N.")

    add_heading(doc, "CHƯƠNG 6. THIẾT KẾ PHẦN MỀM VÀ GIAO THỨC", 1)
    add_heading(doc, "6.1. Môi trường phát triển", 2)
    add_data_table(doc, ["Môi trường", "Board", "Framework", "Thư mục nguồn"], [
        ("esp32", "esp32dev", "Arduino", "src/esp32/"),
        ("uno", "Arduino Uno", "Arduino", "src/arduino/"),
    ], widths=(1.3, 1.5, 1.4, 2.3))
    add_paragraph(doc, "Các thư viện chính gồm PubSubClient, MFRC522, DHT sensor library, LiquidCrystal_I2C, ESP32Servo/Servo và Wire.")

    add_heading(doc, "6.2. Chu trình chương trình ESP32", 2)
    add_paragraph(doc, "Trong hàm setup(), ESP32 khởi tạo LED, I2C master, RFID, DHT11, cảm biến mưa, HC-SR04, cảm biến rung, cảm biến gas, WiFi, LCD/NTP và MQTT. Trong loop(), chương trình gọi lần lượt các module mqttLoop(), rfidServoLoop(), dhtLoop(), rainLoop(), lcdLoop(), radarLoop(), earthquakeLoop(), gasLoop() và updateBuzzer().")
    doc.add_page_break()
    add_heading(doc, "6.3. Giao thức I2C", 2)
    add_data_table(doc, ["Command ID", "Dữ liệu", "Tác dụng"], [
        ("0x01", "90 / 0", "Mở hoặc đóng servo cửa chính."),
        ("0x02", "1", "Kích hoạt LED đỏ/còi báo thẻ RFID sai."),
        ("0x03", "1 / 0", "Bật/tắt đèn phòng khách."),
        ("0x05", "0 / 90", "Điều khiển servo cửa sổ/chan mưa."),
        ("0x06", "1 / 0", "Bật/tắt đèn bếp."),
        ("0x07", "1 / 0", "Bật/tắt đèn phòng ngủ."),
        ("0x08", "1 / 0", "Bật/tắt đèn nhà vệ sinh."),
    ], widths=(1.2, 1.2, 4.1))

    add_heading(doc, "6.4. Topic MQTT", 2)
    add_data_table(doc, ["Topic", "Payload", "Chức năng"], [
        ("home/esp32/led1/set", "ON/OFF", "Điều khiển LED 1."),
        ("home/esp32/servo/set", "OPEN/CLOSE", "Điều khiển cửa chính."),
        ("home/esp32/fan/set", "ON/OFF", "Điều khiển quạt thủ công."),
        ("home/esp32/fan/mode/set", "auto/manual", "Chọn chế độ quạt."),
        ("home/esp32/window/set", "OPEN/CLOSE", "Điều khiển cửa sổ/chan mưa."),
        ("home/esp32/sensor/temperature", "Số thực", "Publish nhiệt độ."),
        ("home/esp32/sensor/humidity", "Số thực", "Publish độ ẩm."),
        ("home/esp32/sensor/rain", "ON/OFF", "Trạng thái mưa."),
        ("home/esp32/sensor/gas", "ON/OFF", "Cảnh báo gas."),
        ("home/esp32/sensor/earthquake", "ON/OFF", "Cảnh báo rung chấn."),
    ], widths=(3.0, 1.3, 2.2))

    add_heading(doc, "CHƯƠNG 7. THIẾT KẾ CÁC CHỨC NĂNG CHÍNH", 1)
    add_heading(doc, "7.1. Khóa cửa RFID", 2)
    add_paragraph(doc, "ESP32 đọc UID từ module MFRC522, chuyển UID sang chuỗi HEX in hoa và so sánh với danh sách thẻ hợp lệ trong config.cpp. Nếu UID hợp lệ, ESP32 gửi lệnh I2C [0x01, 90] để mở cửa và publish trạng thái OPEN; sau 3 giây hệ thống tự gửi [0x01, 0] để đóng cửa. Nếu UID không hợp lệ, ESP32 gửi [0x02, 1] để kích hoạt cảnh báo lỗi và publish log unauthorized.")
    add_heading(doc, "7.2. Điều khiển quạt theo nhiệt độ", 2)
    add_paragraph(doc, "DHT11 được đọc mỗi 5 giây. Dữ liệu nhiệt độ và độ ẩm được publish lên MQTT. Ở chế độ Auto, nếu nhiệt độ lớn hơn 30.0°C, ESP32 bật quạt tại GPIO12; nếu nhiệt độ nhỏ hơn hoặc bằng 30.0°C, quạt tắt. Khi người dùng bấm điều khiển quạt thủ công trên Home Assistant, hệ thống tự chuyển sang Manual để ưu tiên thao tác người dùng.")
    add_heading(doc, "7.3. Cảm biến mưa và cửa sổ/chan mưa", 2)
    add_paragraph(doc, "Cảm biến mưa LM393 được đọc mỗi 500 ms. Khi có nước, chân DO về LOW, ESP32 publish rain = ON và điều khiển servo cửa sổ/chan mưa qua I2C. Khi hết mưa, hệ thống chờ 5 giây rồi mới đưa cơ cấu về trạng thái bình thường để tránh dao động do cảm biến vừa khô hoặc còn đọng nước.")
    add_heading(doc, "7.4. Đèn nhà vệ sinh tự động", 2)
    add_paragraph(doc, "HC-SR04 phát xung siêu âm và đo thời gian phản hồi để tính khoảng cách. Nếu khoảng cách nhỏ hơn 10 cm, hệ thống xem như có người và bật đèn nhà vệ sinh. Khi không còn phát hiện người, hệ thống chờ 5 giây rồi tắt đèn, đồng thời publish trạng thái lên MQTT.")
    add_heading(doc, "7.5. Cảnh báo gas và động đất", 2)
    add_paragraph(doc, "Cảm biến gas tại GPIO4 và cảm biến rung tại GPIO32 được xử lý độc lập, nhưng cùng sử dụng buzzer GPIO33. Hàm updateBuzzer() bật còi khi có ít nhất một trong hai trạng thái nguy hiểm: gas leaking hoặc quaking. Thiết kế này đơn giản, dễ hiểu và phù hợp với mô hình cảnh báo tổng hợp.")
    add_heading(doc, "7.6. LCD và thời gian NTP", 2)
    add_paragraph(doc, "LCD I2C 16x2 được khởi tạo sau khi ESP32 kết nối WiFi để có thể đồng bộ thời gian từ NTP server pool.ntp.org. Dòng đầu hiển thị ngày giờ theo múi giờ Việt Nam UTC+7, dòng thứ hai hiển thị nhiệt độ và độ ẩm.")

    add_heading(doc, "CHƯƠNG 8. MÔ PHỎNG, CÀI ĐẶT VÀ VẬN HÀNH", 1)
    add_heading(doc, "8.1. Chạy Mosquitto và Home Assistant", 2)
    add_paragraph(doc, "Server cục bộ được triển khai bằng Docker Compose. Mosquitto mở cổng 1883 để ESP32 kết nối MQTT, Home Assistant mở cổng 8123 để người dùng truy cập dashboard. Trong môi trường mô phỏng, Mosquitto cho phép anonymous để giảm độ phức tạp cấu hình.")
    add_heading(doc, "8.2. Biên dịch và nạp chương trình", 2)
    add_data_table(doc, ["Tác vụ", "Lệnh"], [
        ("Chạy server", "docker compose up -d"),
        ("Build ESP32", "pio run -e esp32"),
        ("Upload ESP32", "pio run -e esp32 -t upload"),
        ("Build Arduino Uno", "pio run -e uno"),
        ("Upload Arduino Uno", "pio run -e uno -t upload"),
    ], widths=(1.8, 4.7))
    add_heading(doc, "8.3. Kịch bản vận hành", 2)
    add_numbered(doc, [
        "Điều khiển đèn trên Home Assistant và quan sát LED tương ứng thay đổi.",
        "Quét thẻ RFID hợp lệ để mở cửa; sau 3 giây cửa tự đóng.",
        "Quét thẻ không hợp lệ để kiểm tra LED đỏ/còi cảnh báo.",
        "Bật chế độ quạt tự động, làm tăng nhiệt độ quanh DHT11 và quan sát quạt bật khi vượt 30.0°C.",
        "Nhỏ nước lên cảm biến mưa để kiểm tra cơ cấu che mưa/cửa sổ.",
        "Đưa vật thể lại gần HC-SR04 để kiểm tra đèn nhà vệ sinh tự động.",
        "Kích hoạt cảm biến rung hoặc gas để kiểm tra cảnh báo an toàn.",
    ])

    add_heading(doc, "CHƯƠNG 9. KIỂM THỬ VÀ ĐÁNH GIÁ KẾT QUẢ", 1)
    add_data_table(doc, ["STT", "Chức năng", "Kết quả mong đợi", "Đánh giá"], [
        (1, "Kết nối WiFi", "ESP32 nhận IP trong mạng LAN.", "Đạt"),
        (2, "Kết nối MQTT", "ESP32 subscribe topic và đồng bộ trạng thái.", "Đạt"),
        (3, "Điều khiển đèn", "LED/đèn mô phỏng đổi trạng thái theo Home Assistant.", "Đạt"),
        (4, "RFID hợp lệ", "Servo mở cửa, LED xanh sáng, publish OPEN.", "Đạt"),
        (5, "RFID sai", "LED đỏ/cảnh báo lỗi, publish unauthorized.", "Đạt"),
        (6, "DHT11 và quạt", "Publish nhiệt độ/độ ẩm, quạt bật khi > 30.0°C.", "Đạt"),
        (7, "Cảm biến mưa", "Rain ON và servo chan mưa hoạt động.", "Đạt"),
        (8, "HC-SR04", "Đèn nhà vệ sinh bật khi phát hiện người.", "Đạt"),
        (9, "Gas/rung", "Binary sensor ON và buzzer kích hoạt.", "Đạt"),
        (10, "LCD", "Hiển thị ngày giờ, nhiệt độ và độ ẩm.", "Đạt"),
    ], widths=(0.55, 1.6, 3.35, 1.0))
    add_heading(doc, "9.1. Ưu điểm", 2)
    add_bullets(doc, [
        "Kiến trúc rõ ràng, tách module theo từng chức năng.",
        "Sử dụng MQTT và Home Assistant giúp hệ thống dễ giám sát và mở rộng.",
        "Kết hợp ESP32 và Arduino Uno giúp mô phỏng kiến trúc nhiều vi điều khiển.",
        "Có cả điều khiển thủ công và tự động hóa theo dữ liệu cảm biến.",
    ])
    add_heading(doc, "9.2. Hạn chế", 2)
    add_bullets(doc, [
        "MQTT Broker đang cho phép anonymous để tiện mô phỏng, chưa phù hợp triển khai thực tế.",
        "Danh sách UID RFID còn cố định trong code, chưa có giao diện quản lý thẻ.",
        "DHT11 và cảm biến gas cần hiệu chuẩn nếu dùng trong môi trường thật.",
        "Mô hình còn ở mức breadboard, chưa có PCB, vỏ bảo vệ và nguồn công suất ổn định.",
    ])

    add_heading(doc, "CHƯƠNG 10. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "10.1. Kết luận", 2)
    add_paragraph(doc, "Đồ án đã thiết kế và mô phỏng thành công hệ thống nhà thông minh sử dụng ESP32, Arduino Uno, MQTT và Home Assistant. Hệ thống đáp ứng các chức năng quan trọng như điều khiển đèn, mở cửa bằng RFID, đo nhiệt độ/độ ẩm, điều khiển quạt tự động, xử lý khi trời mưa, bật đèn nhà vệ sinh theo phát hiện người và cảnh báo an toàn khi có gas hoặc rung chấn.")
    add_paragraph(doc, "Thông qua đề tài, người thực hiện nắm được cách kết hợp phần cứng nhúng, giao thức nội bộ I2C, giao thức mạng MQTT và nền tảng dashboard Home Assistant để tạo thành một hệ thống IoT hoàn chỉnh.")
    add_heading(doc, "10.2. Hướng phát triển", 2)
    add_bullets(doc, [
        "Bổ sung username/password hoặc TLS cho MQTT.",
        "Xây dựng giao diện thêm/xóa thẻ RFID từ Home Assistant.",
        "Lưu lịch sử dữ liệu vào InfluxDB và hiển thị biểu đồ bằng Grafana.",
        "Gửi cảnh báo qua Telegram, email hoặc push notification.",
        "Nâng cấp DHT11 lên DHT22/SHT31 để tăng độ chính xác.",
        "Thiết kế PCB, vỏ bảo vệ và nguồn riêng cho servo, quạt.",
    ])

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_numbered(doc, [
        "Espressif Systems, ESP32 Technical Reference Manual.",
        "Arduino Documentation, Arduino Uno Rev3.",
        "MQTT Version 3.1.1 Specification.",
        "Home Assistant Documentation - MQTT Integration.",
        "Eclipse Mosquitto Documentation.",
        "PlatformIO Documentation.",
        "Datasheet DHT11 và MFRC522.",
        "Tài liệu nội bộ trong thư mục .doc của đồ án.",
    ])

    add_heading(doc, "PHỤ LỤC. THÔNG SỐ QUAN TRỌNG", 1)
    add_data_table(doc, ["Thông số", "Giá trị"], [
        ("Baudrate ESP32", "115200"),
        ("Baudrate Arduino Uno", "9600"),
        ("MQTT Port", "1883"),
        ("Home Assistant Port", "8123"),
        ("I2C Slave Address", "0x08"),
        ("ESP32 SDA/SCL", "GPIO21 / GPIO22"),
        ("DHT interval", "5000 ms"),
        ("Rain/Radar interval", "500 ms"),
        ("LCD interval", "1000 ms"),
        ("RFID auto close", "3000 ms"),
        ("Rain dry delay", "5000 ms"),
        ("Bathroom light off delay", "5000 ms"),
        ("Temperature threshold", "30.0°C"),
        ("Radar detection threshold", "< 10 cm"),
        ("Buzzer pin", "GPIO33"),
    ], widths=(2.5, 4.0))

    doc.save(OUT)


if __name__ == "__main__":
    build()
